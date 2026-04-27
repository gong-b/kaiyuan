import os
import re
import tempfile
from docx import Document
from pdf2docx import Converter

class FileParser:
    @staticmethod
    def parse(path):
        ext = str(path).lower()
        if ext.endswith('.pdf'):
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                tmp_docx_path = tmp_docx.name
            
            try:
                cv = Converter(path)
                cv.convert(tmp_docx_path, start=0, end=1)
                cv.close()
                res = FileParser._parse_docx(tmp_docx_path)
                if os.path.exists(tmp_docx_path):
                    os.remove(tmp_docx_path)
                return res
            except Exception as e:
                print(f"PDF 转换失败: {e}")
                return {"is_supported": False, "reason_length": 0, "name": "转换失败", "sid": None, "apply_class": "PDF解析异常", "contact": ""}
                
        return FileParser._parse_docx(path)

    @staticmethod
    def _parse_docx(path):
        res = {
            "is_supported": False,
            "reason_length": 0,
            "name": "未知",
            "sid": None,
            "apply_class": "",
            "contact": ""
        }
        try:
            doc = Document(path)
            
            # 提取班级
            for para in doc.paragraphs:
                text = para.text.replace(" ", "")
                match = re.search(r"(.+?班)报名申请表", text)
                if match:
                    res["apply_class"] = match.group(1)
                    break

            # 表格解析（修复版：支持两种申请理由格式）
            if doc.tables:
                table = doc.tables[0]
                # 先把所有单元格内容按顺序存起来，方便后续处理
                all_cells = []
                for row in table.rows:
                    for cell in row.cells:
                        all_cells.append(cell)

                # 遍历所有单元格，解析信息
                for i, cell in enumerate(all_cells):
                    cell_text = cell.text.strip()
                    # 姓名
                    if cell_text == "姓名" and i + 1 < len(all_cells):
                        res["name"] = all_cells[i+1].text.strip()
                    # 学号
                    elif cell_text == "学号" and i + 1 < len(all_cells):
                        sid_text = all_cells[i+1].text.strip()
                        res["sid"] = "".join(filter(str.isdigit, sid_text))
                    # 联系方式
                    elif any(key in cell_text for key in ["联系方式", "电话", "手机", "联系电话"]):
                        if i + 1 < len(all_cells):
                            contact_text = all_cells[i+1].text.strip()
                            contact = re.sub(r"[^\d\- ]", "", contact_text)
                            res["contact"] = contact.strip()
                    # 资助对象
                    elif "资助对象" in cell_text:
                        res["is_supported"] = "是" in cell_text

                    # ====================== 终极修复：申请理由解析 ======================
                    elif "申请理由" in cell_text:
                        # 你指定的标准：先按 paragraphs 读
                        reason_paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                        reason_text = "\n".join(reason_paragraphs)
                        reason_text = re.sub(r"\s+", "", reason_text)
                        count = len(reason_text.replace(" ", ""))

                        # 🔥 关键修复：如果 paragraphs 读出来为0，强制用 cell.text 兜底（100%能读到）
                        if count == 0:
                            full_text = cell.text.strip()
                            full_text = re.sub(r"申请理由.*?[:：]", "", full_text)
                            full_text = re.sub(r"\s+", "", full_text)
                            count = len(full_text)

                        # 再检查下一格（兼容你的表格格式）
                        if count == 0 and i + 1 < len(all_cells):
                            next_cell = all_cells[i+1]
                            next_text = next_cell.text.strip()
                            next_text = re.sub(r"\s+", "", next_text)
                            count = len(next_text)

                        res["reason_length"] = count
                    # ====================================================================

        except Exception as e:
            print(f"解析异常: {e}")
            return res
        return res
