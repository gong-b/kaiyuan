from docx import Document
import re

class DocxParser:
    @staticmethod
    def parse(path):
    res = {
        "is_supported": False, 
        "reason_length": 0, 
        "name": "未知", 
        "sid": None,
        "apply_class": ""  # 新增：报名班级
    }
    try:
        doc = Document(path)
        if not doc.tables:
            return res
        
        table = doc.tables[0]
        for row in table.rows:
            cells = row.cells
            for i, c in enumerate(cells):
                t = c.text.strip()
                
                # 原有逻辑：提取姓名、学号、资助对象、申请理由...
                if t == "姓名" and i + 1 < len(cells):
                    res["name"] = cells[i+1].text.strip()
                if t == "学号" and i + 1 < len(cells):
                    raw_sid = cells[i+1].text.strip()
                    res["sid"] = "".join(filter(str.isdigit, raw_sid))
                if "是否为学生资助对象" in t and i + 1 < len(cells):
                    v = cells[i + 1].text.strip()
                    res["is_supported"] = ("是" in v) and ("不是" not in v)
                if "申请理由" in t:
                    import re
                    pattern = re.compile(r"申请理由\s*[:：]\s*|申请理由\s*（.*?）\s*[:：]\s*")
                    content = pattern.sub("", c.text).strip()
                    clean_content = re.sub(r"\s+", "", content)
                    res["reason_length"] = len(clean_content)
                
                # 新增：提取报名班级（匹配“报名班级”“申请班级”等字段）
                if "报名班级" in t and i + 1 < len(cells):
                    res["apply_class"] = cells[i+1].text.strip()
                elif "申请班级" in t and i + 1 < len(cells):
                    res["apply_class"] = cells[i+1].text.strip()
    except Exception as e:
        print(f"解析附件出错: {e}")
    return res
