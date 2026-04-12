import logging
import re
import os
from datetime import datetime, timezone
from email.header import decode_header
from email.message import Message
from email.utils import parsedate_to_datetime
from pathlib import Path

# ---------------------- 配置项（独立配置，避免依赖外部config.py） ----------------------
# 基础目录
BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(exist_ok=True, parents=True)

# 文件路径配置
NEW_HONGJI_FILE = DATA_DIR / "2024-2025学年秋冬学期新鸿基推荐学生名单.xlsx"
LAST_YEAR_FILE = DATA_DIR / "24秋冬学期开源课堂人员名单.xlsx"
BLACKLIST_FILE = DATA_DIR / "blacklist.xlsx"
ADMITTED_FILE = DATA_DIR / "admitted_students.xlsx"
REJECTED_FILE = DATA_DIR / "rejected_students.xlsx"

# 邮箱配置（从环境变量读取）
IMAP_HOST = os.environ.get("IMAP_HOST", "imap.zju.edu.cn")
IMAP_PORT = int(os.environ.get("IMAP_PORT", 993))
EMAIL_USER = os.environ.get("EMAIL_USER", "")
EMAIL_PASSWORD = os.environ.get("EMAIL_PASSWORD", "")

# 录取配置
ADMISSION_QUOTA = int(os.environ.get("ADMISSION_QUOTA", 25))

# ---------------------- 日志配置 ----------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(module)s:%(lineno)d - %(message)s",
    handlers=[
        logging.FileHandler(DATA_DIR / "processing.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ---------------------- 正则表达式 ----------------------
SUBJECT_PATTERN = re.compile(
    r"^\s*[()（）\[\]【】\{\}｛｝]*([\u4e00-\u9fa5]{2,4})\s*[+＋-—\s]*(\d{8,12})\s*[+＋-—\s]*书法班报名申请[()（）\[\]【】\{\}｛｝]*\s*$",
    re.IGNORECASE
)

# ---------------------- 邮箱客户端类（独立实现，无外部依赖） ----------------------
import imaplib
import ssl
from typing import Generator, Tuple, Optional

class SecureIMAPClient:
    """IMAP客户端（增强异常处理）"""
    def __init__(self) -> None:
        self.host = IMAP_HOST
        self.port = IMAP_PORT
        self.user = EMAIL_USER
        self.password = EMAIL_PASSWORD
        self.mailbox = "INBOX"
        self.conn: Optional[imaplib.IMAP4_SSL] = None

    def __enter__(self) -> "SecureIMAPClient":
        """上下文进入：精准捕获IMAP连接/登录异常"""
        # 前置校验（非空）
        required_config = [self.user, self.password, self.host]
        if not all(required_config):
            raise ValueError(
                "IMAP配置不完整：\n"
                f"- 邮箱账号: {'已配置' if self.user else '缺失'}\n"
                f"- 客户端密码: {'已配置' if self.password else '缺失'}\n"
                f"- IMAP主机: {self.host}"
            )
        
        # 构建SSL上下文
        context: ssl.SSLContext = ssl.create_default_context()
        context.check_hostname = False
        context.verify_mode = ssl.CERT_NONE
        
        try:
            # 连接IMAP服务器（捕获连接异常）
            self.conn = imaplib.IMAP4_SSL(
                self.host, 
                self.port, 
                ssl_context=context,
                timeout=15  # 连接超时15秒
            )
            logger.debug(f"SSL连接建立成功: {self.host}:{self.port}")
            
            # 登录（捕获登录异常）
            try:
                self.conn.login(self.user.encode('utf-8'), self.password.encode('utf-8'))
            except imaplib.IMAP4.error as e:
                error_msg = str(e).strip()
                # 针对性提示
                if "authentication failed" in error_msg.lower():
                    raise RuntimeError(
                        "邮箱登录失败！请检查：\n"
                        "1. 客户端专用密码是否正确（非登录密码）\n"
                        "2. 浙大邮箱是否开启IMAP/SMTP服务\n"
                        "3. 账号是否有IMAP访问权限"
                    ) from e
                else:
                    raise RuntimeError(f"IMAP登录错误: {error_msg}") from e
            
            # 选择邮箱文件夹（捕获文件夹异常）
            try:
                status, _ = self.conn.select(self.mailbox, readonly=True)
                if status != "OK":
                    raise RuntimeError(f"选择邮箱文件夹失败: {self.mailbox}")
            except imaplib.IMAP4.error as e:
                raise RuntimeError(f"访问邮箱文件夹失败: {str(e)}") from e
            
            logger.info(f"成功登录邮箱: {self.user}")
            return self
        
        # 捕获连接类异常
        except TimeoutError:
            raise RuntimeError(f"连接IMAP服务器超时（{self.host}:{self.port}），请检查网络/服务器状态")
        except ConnectionRefusedError:
            raise RuntimeError(f"IMAP服务器拒绝连接（{self.host}:{self.port}），请检查主机/端口配置")
        except Exception as e:
            raise RuntimeError(f"邮箱连接初始化失败: {str(e)}") from e

    def __exit__(self, exc_type, exc_value, traceback) -> None:
        """上下文退出：安全关闭连接（忽略退出时的小错误）"""
        if self.conn:
            try:
                if self.conn.state == "SELECTED":
                    self.conn.close()
                self.conn.logout()
                logger.info("IMAP连接已安全关闭")
            except Exception as e:
                # 退出时的错误不影响主流程，仅日志记录
                logger.warning(f"关闭IMAP连接时警告: {str(e)}")
            finally:
                self.conn = None

    def fetch_emails(self) -> Generator[Tuple[str, Message], None, None]:
        """获取邮件：分步骤捕获异常，出错跳过单封邮件"""
        if self.conn is None:
            logger.error("IMAP连接未初始化，无法获取邮件")
            return

        # 1. 解析日期范围（捕获日期格式异常）
        try:
            start_date_str = os.environ.get("START_DATE", "01-Mar-2025")
            end_date_str = os.environ.get("END_DATE", datetime.now().strftime("%d-%b-%Y"))
            # 验证日期格式（IMAP要求：dd-Mon-yyyy）
            datetime.strptime(start_date_str, "%d-%b-%Y")
            datetime.strptime(end_date_str, "%d-%b-%Y")
        except ValueError as e:
            logger.error(f"日期格式错误（需dd-Mon-yyyy，如01-Mar-2025）: {e}")
            return

        # 2. 搜索邮件（捕获搜索异常）
        try:
            search_criteria = f'(SINCE "{start_date_str}" BEFORE "{end_date_str}")'
            logger.info(f"搜索邮件条件: {search_criteria}")
            status, data = self.conn.uid('SEARCH', None, search_criteria)
            
            if status != 'OK':
                error_msg = data[0].decode('utf-8', errors='replace') if data else "未知错误"
                logger.error(f"邮件搜索失败: {error_msg}")
                return

            uids = data[0].split() if data and data[0] else []
            if not uids:
                logger.info("未找到符合日期条件的邮件")
                return
            logger.info(f"找到{len(uids)}封待处理邮件")
        
        except imaplib.IMAP4.error as e:
            logger.error(f"IMAP搜索邮件异常: {str(e)}", exc_info=True)
            return
        except Exception as e:
            logger.error(f"邮件搜索流程异常: {str(e)}", exc_info=True)
            return

        # 3. 遍历邮件（单封邮件出错不终止，仅跳过）
        for uid_bytes in uids:
            uid = uid_bytes.decode('utf-8').strip()
            if not uid:
                logger.warning("空UID，跳过")
                continue
                
            try:
                # 获取邮件原始内容
                status, msg_data = self.conn.uid('FETCH', uid, '(RFC822)')
                if status != 'OK':
                    logger.error(f"获取邮件{uid}失败: {msg_data}")
                    continue

                # 解析邮件数据
                raw_email = None
                for response_part in msg_data:
                    if isinstance(response_part, tuple) and len(response_part) >= 2:
                        raw_email = response_part[1]
                        break
                
                if not isinstance(raw_email, bytes):
                    logger.error(f"邮件{uid}数据格式异常（非字节类型）: {type(raw_email)}")
                    continue

                # 解析为Message对象
                msg = message_from_bytes(raw_email)
                
                # 过滤书法班邮件
                subject = self._get_msg_subject(msg)
                if "书法班" in subject:
                    yield uid, msg
                else:
                    logger.debug(f"邮件{uid}非书法班报名（主题：{subject[:50]}...），跳过")
                    
            except Exception as e:
                # 单封邮件出错，记录日志后继续处理下一封
                logger.error(f"处理邮件{uid}失败（已跳过）: {str(e)}", exc_info=True)
                continue

    @staticmethod
    def _get_msg_subject(msg: Message) -> str:
        """获取主题：捕获解码异常"""
        subject = msg.get("Subject", "")
        decoded_parts = []
        for part, charset in decode_header(subject):
            try:
                if isinstance(part, bytes):
                    # 优先尝试常见编码
                    for encoding in [charset, 'utf-8', 'gb18030', 'big5', 'gbk']:
                        if not encoding:
                            continue
                        try:
                            decoded = part.decode(encoding)
                            break
                        except:
                            continue
                    else:
                        decoded = part.decode('utf-8', errors='replace')
                else:
                    decoded = str(part)
                decoded_parts.append(decoded)
            except Exception as e:
                logger.warning(f"主题片段解码失败: {e}")
                decoded_parts.append("[解码失败]")
        return "".join(decoded_parts)

# ---------------------- 邮件附件处理器 ----------------------
import os
from email.message import Message
from pathlib import Path

class EmailProcessor:
    """邮件附件处理器"""
    def __init__(self):
        self.attach_dir = DATA_DIR / "attachments"
        self.attach_dir.mkdir(exist_ok=True, parents=True)

    def save_attachments(self, msg: Message, student_id: str, name: str) -> list[Path]:
        """保存邮件附件，返回附件路径列表"""
        attachments = []
        try:
            for part in msg.walk():
                # 跳过非附件部分
                if part.get_content_maintype() == 'multipart':
                    continue
                if part.get('Content-Disposition') is None:
                    continue

                # 获取附件文件名
                filename = part.get_filename()
                if not filename:
                    continue

                # 解码文件名
                decoded_filename = self._decode_filename(filename)
                # 重命名文件（避免重复）
                file_ext = Path(decoded_filename).suffix
                safe_filename = f"{student_id}_{name}_{datetime.now().strftime('%Y%m%d%H%M%S')}{file_ext}"
                file_path = self.attach_dir / safe_filename

                # 保存附件
                try:
                    with open(file_path, 'wb') as f:
                        f.write(part.get_payload(decode=True))
                    attachments.append(file_path)
                    logger.debug(f"保存附件成功: {file_path}")
                except Exception as e:
                    logger.error(f"保存附件{decoded_filename}失败: {str(e)}")
                    continue

            return attachments
        except Exception as e:
            logger.error(f"处理附件失败: {str(e)}", exc_info=True)
            return attachments

    @staticmethod
    def _decode_filename(filename: str) -> str:
        """解码附件文件名"""
        try:
            decoded_parts = decode_header(filename)
            filename_parts = []
            for part, charset in decoded_parts:
                if isinstance(part, bytes):
                    if charset:
                        filename_parts.append(part.decode(charset, errors='replace'))
                    else:
                        filename_parts.append(part.decode('utf-8', errors='replace'))
                else:
                    filename_parts.append(str(part))
            return ''.join(filename_parts)
        except Exception as e:
            logger.warning(f"解码文件名失败: {str(e)}")
            return f"unknown_{datetime.now().strftime('%Y%m%d%H%M%S')}"

# ---------------------- DOCX解析器 ----------------------
from docx import Document
from typing import Dict, Any

def parse_docx(filepath: str | Path) -> Dict[str, Any]:
    """解析DOCX：精准捕获docx相关异常"""
    result = {"is_supported": False, "reason_length": 0}
    filepath = Path(filepath)

    if not filepath.exists():
        logger.warning(f"DOCX文件不存在: {filepath}")
        return result
    if filepath.stat().st_size == 0:
        logger.warning(f"DOCX文件为空: {filepath}")
        return result

    try:
        # 捕获docx打开异常
        try:
            doc = Document(filepath)
        except Exception as e:
            logger.error(f"打开DOCX文件失败（可能文件损坏）: {filepath} - {e}", exc_info=True)
            return result

        if not doc.tables:
            logger.warning(f"DOCX文件无表格: {filepath}")
            return result

        # 遍历表格（捕获索引异常）
        support_flag = False
        reason_text = ""
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                try:
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if not cell_text:
                            continue

                        # 匹配资助对象
                        if any(kw in cell_text for kw in ["是否为学生资助对象", "资助对象", "贫困生"]):
                            if cell_idx + 1 < len(row.cells):
                                next_cell = row.cells[cell_idx+1].text.strip()
                                support_flag = any(yes in next_cell for yes in ["是", "√"]) and not any(no in next_cell for no in ["否", "×"])

                        # 匹配申请理由
                        if any(kw in cell_text for kw in ["申请理由", "申请原因"]):
                            reason_parts = [cell_text]
                            # 避免索引越界
                            if cell_idx + 1 < len(row.cells):
                                reason_parts.append(row.cells[cell_idx+1].text.strip())
                            # 避免行索引越界
                            if row_idx + 1 < len(table.rows):
                                next_row = table.rows[row_idx+1]
                                reason_parts.extend([c.text.strip() for c in next_row.cells])
                            reason_text = "".join(reason_parts)
                except IndexError as e:
                    logger.debug(f"表格{table_idx}行{row_idx}单元格索引异常: {e}")
                    continue
                except Exception as e:
                    logger.debug(f"解析表格{table_idx}行{row_idx}失败: {e}")
                    continue

        # 清理理由文本
        reason_text = re.sub(r"\s+", "", reason_text)
        reason_text = re.sub(r"[^\u4e00-\u9fa5]", "", reason_text)
        result["is_supported"] = support_flag
        result["reason_length"] = len(reason_text)
        return result

    except Exception as e:
        logger.error(f"解析DOCX未预期异常: {filepath} - {e}", exc_info=True)
        return result

# ---------------------- Excel处理函数 ----------------------
import pandas as pd
from typing import Set

def read_student_list(file_path: str | Path) -> Set[str]:
    """
    读取学生名单，返回学号集合（增强容错，精准捕获Excel相关异常）
    :param file_path: Excel文件路径
    :return: 有效学号集合
    """
    student_ids = set()
    try:
        # 路径标准化
        file_path = Path(file_path)
        
        # 前置校验：文件存在性+非空
        if not file_path.exists():
            logger.warning(f"Excel文件不存在: {file_path}")
            return student_ids
        if file_path.stat().st_size == 0:
            logger.warning(f"Excel文件为空（大小0字节）: {file_path}")
            return student_ids
        
        # 尝试打开Excel文件（捕获特定异常）
        try:
            xl_file = pd.ExcelFile(file_path, engine="openpyxl")
        except pd.errors.EmptyDataError:
            logger.warning(f"Excel文件无任何数据: {file_path}")
            return student_ids
        except ImportError:
            logger.error(f"缺少openpyxl依赖！请执行: pip install openpyxl")
            return student_ids
        except Exception as e:
            logger.error(f"打开Excel文件失败（可能文件损坏/格式错误）: {file_path} - {str(e)}", exc_info=True)
            return student_ids
        
        # 自动识别包含"学号"的Sheet（遍历所有Sheet）
        target_sheet = None
        for sheet_name in xl_file.sheet_names:
            try:
                # 仅读取前10行检测列名，提升效率
                df_temp = pd.read_excel(xl_file, sheet_name=sheet_name, nrows=10)
                # 匹配包含"学号"的列（不区分大小写/全半角）
                if any("学号" in str(col).lower().replace(" ", "") for col in df_temp.columns):
                    target_sheet = sheet_name
                    logger.debug(f"找到包含学号列的Sheet: {sheet_name}")
                    break
            except Exception as e:
                logger.debug(f"检查Sheet [{sheet_name}] 失败（跳过）: {str(e)}")
                continue
        
        if not target_sheet:
            logger.warning(f"文件{file_path}中未找到包含'学号'列的Sheet")
            return student_ids
        
        # 读取目标Sheet数据（捕获解析异常）
        try:
            df = pd.read_excel(
                xl_file,
                sheet_name=target_sheet,
                engine="openpyxl",
                dtype=str  # 强制所有列按字符串读取，避免学号被解析为数字
            )
        except Exception as e:
            logger.error(f"读取Sheet [{target_sheet}] 失败: {file_path} - {str(e)}", exc_info=True)
            return student_ids
        
        # 查找学号列（兼容不同列名：学号、学生学号、ID等）
        id_cols = [col for col in df.columns if "学号" in str(col).lower().replace(" ", "")]
        if not id_cols:
            logger.warning(f"Sheet [{target_sheet}] 中未找到'学号'相关列")
            return student_ids
        
        # 提取并清洗学号（去重、去空、去非数字）
        id_series = df[id_cols[0]].astype(str).dropna().str.strip()
        # 过滤规则：非空 + 不是"nan" + 纯数字
        student_ids = set(
            sid for sid in id_series 
            if sid and sid != "nan" and sid.replace(" ", "").isdigit()
        )
        
        logger.info(f"读取学生名单成功: {file_path} - Sheet[{target_sheet}] - 有效学号数: {len(student_ids)}")
        return student_ids
        
    except Exception as e:
        logger.error(f"读取学生名单未预期异常: {file_path} - {str(e)}", exc_info=True)
        return student_ids

def save_results(admitted: list[dict[str, Any]], rejected: list[dict[str, Any]]):
    """
    保存录取/拒绝结果（增强容错，确保文件可写、列完整）
    :param admitted: 录取名单列表
    :param rejected: 拒绝名单列表
    """
    try:
        # 确保输出目录存在
        for file_path in [ADMITTED_FILE, REJECTED_FILE]:
            file_path.parent.mkdir(exist_ok=True, parents=True)
        
        # ---------------------- 保存录取名单 ----------------------
        if admitted:
            # 转换为DataFrame，确保核心列存在
            df_admitted = pd.DataFrame(admitted)
            # 补全缺失列（避免KeyError）
            for col in ["学号", "姓名"]:
                if col not in df_admitted.columns:
                    df_admitted[col] = ""
            # 补充备注列（可选）
            if "备注" not in df_admitted.columns:
                df_admitted["备注"] = ""
            # 列顺序标准化
            df_admitted = df_admitted[["学号", "姓名", "备注"]]
            
            # 去重（按学号，保留第一条）
            df_admitted = df_admitted.drop_duplicates(subset=["学号"], keep="first")
            
            # 保存文件（覆盖已有文件，编码UTF-8）
            try:
                df_admitted.to_excel(
                    ADMITTED_FILE,
                    sheet_name="录取名单",
                    index=False,
                    engine="openpyxl",
                    encoding="utf-8"
                )
                logger.info(f"录取名单已保存: {ADMITTED_FILE} - 共{len(df_admitted)}条记录")
            except PermissionError:
                logger.error(f"保存录取名单失败：无写入权限 - {ADMITTED_FILE}")
                raise RuntimeError(f"无法写入文件（权限不足）: {ADMITTED_FILE}")
            except Exception as e:
                logger.error(f"保存录取名单失败: {ADMITTED_FILE} - {str(e)}", exc_info=True)
                raise RuntimeError(f"保存录取名单失败: {str(e)}")
        
        # ---------------------- 保存拒绝名单 ----------------------
        if rejected:
            # 转换为DataFrame，确保核心列存在
            df_rejected = pd.DataFrame(rejected)
            # 补全缺失列
            for col in ["学号", "姓名", "原因"]:
                if col not in df_rejected.columns:
                    df_rejected[col] = ""
            # 补充原主题列（可选）
            if "原主题" not in df_rejected.columns:
                df_rejected["原主题"] = ""
            # 列顺序标准化
            df_rejected = df_rejected[["学号", "姓名", "原主题", "原因"]]
            
            # 去重（按学号+原因，保留第一条）
            df_rejected = df_rejected.drop_duplicates(subset=["学号", "原因"], keep="first")
            
            # 保存文件
            try:
                df_rejected.to_excel(
                    REJECTED_FILE,
                    sheet_name="拒绝名单",
                    index=False,
                    engine="openpyxl",
                    encoding="utf-8"
                )
                logger.info(f"拒绝名单已保存: {REJECTED_FILE} - 共{len(df_rejected)}条记录")
            except PermissionError:
                logger.error(f"保存拒绝名单失败：无写入权限 - {REJECTED_FILE}")
                raise RuntimeError(f"无法写入文件（权限不足）: {REJECTED_FILE}")
            except Exception as e:
                logger.error(f"保存拒绝名单失败: {REJECTED_FILE} - {str(e)}", exc_info=True)
                raise RuntimeError(f"保存拒绝名单失败: {str(e)}")
        
        # 无数据时的提示
        if not admitted and not rejected:
            logger.warning("录取/拒绝名单均为空，未生成任何Excel文件")
        
    except Exception as e:
        logger.error(f"保存结果流程未预期异常: {str(e)}", exc_info=True)
        raise

# ---------------------- 主题解析函数 ----------------------
def parse_subject_pattern(subject: str) -> tuple[str, str] | tuple[None, None]:
    """解析主题：捕获正则匹配异常"""
    if not subject:
        return None, None
    try:
        clean_subject = re.sub(r"\s+", "", subject)
        match = SUBJECT_PATTERN.match(clean_subject)
        if match:
            return match.group(1).strip(), match.group(2).strip()
        return None, None
    except Exception as e:
        logger.error(f"主题正则匹配异常（主题：{subject[:50]}）: {e}", exc_info=True)
        return None, None

# ---------------------- 主执行函数 ----------------------
def main():
    """主流程：分层捕获异常，出错有兜底"""
    # 初始化结果容器（即使中间出错，也能保存已有结果）
    admitted: list[dict[str, str]] = []
    rejected: list[dict[str, str]] = []
    candidates: list[tuple[str, str, datetime]] = []

    try:
        logger.info("="*50 + " 开始执行书法班筛选流程 " + "="*50)
        
        # 1. 初始化处理器（捕获初始化异常）
        try:
            email_processor = EmailProcessor()
        except Exception as e:
            logger.error(f"邮件处理器初始化失败: {e}", exc_info=True)
            raise RuntimeError("邮件处理器初始化失败，请检查依赖配置") from e

        # 2. 读取基础名单（捕获Excel读取异常，单独处理每个文件）
        logger.info("读取基础学生名单...")
        list_config = {
            "新鸿基推荐名单": (NEW_HONGJI_FILE, True),  # 必选
            "去年录取名单": (LAST_YEAR_FILE, True),     # 必选
            "黑名单": (BLACKLIST_FILE, False)           # 可选
        }
        list_data = {}
        
        for list_name, (file_path, required) in list_config.items():
            try:
                data = read_student_list(str(file_path))
                list_data[list_name] = data
                # 必选文件为空则终止
                if required and not data:
                    raise RuntimeError(f"{list_name}为空或文件解析失败: {file_path}")
                logger.info(f"{list_name}读取完成: {len(data)}个有效学号")
            except Exception as e:
                if required:
                    raise RuntimeError(f"读取{list_name}失败（必选文件）: {e}") from e
                else:
                    logger.warning(f"读取{list_name}失败（可选文件，已跳过）: {e}")
                    list_data[list_name] = set()

        new_hongji = list_data["新鸿基推荐名单"]
        last_year = list_data["去年录取名单"]
        blacklist = list_data["黑名单"]

        # 3. 解析日期范围（捕获日期解析异常）
        try:
            start_date_str = os.environ.get("START_DATE", "01-Mar-2025")
            end_date_str = os.environ.get("END_DATE", datetime.now().strftime("%d-%b-%Y"))
            start_date = datetime.strptime(start_date_str, "%d-%b-%Y").replace(tzinfo=timezone.utc)
            end_date = datetime.strptime(end_date_str, "%d-%b-%Y").replace(tzinfo=timezone.utc)
            logger.info(f"处理日期范围: {start_date_str} 至 {end_date_str}")
        except ValueError as e:
            raise RuntimeError(f"日期解析失败（格式需为dd-Mon-yyyy，如01-Mar-2025）: {e}") from e

        # 4. 处理邮件（核心逻辑，单封邮件出错不终止）
        logger.info("开始处理邮件...")
        email_count = 0
        error_count = 0
        
        with SecureIMAPClient() as client:
            for uid, msg in client.fetch_emails():
                email_count += 1
                try:
                    # 解析接收时间
                    recv_date = None
                    date_str = msg.get("Date")
                    if date_str:
                        recv_date = parsedate_to_datetime(date_str)
                        if recv_date.tzinfo is None:
                            recv_date = recv_date.replace(tzinfo=timezone.utc)
                        else:
                            recv_date = recv_date.astimezone(timezone.utc)
                    
                    # 日期过滤
                    if not recv_date or not (start_date <= recv_date <= end_date):
                        logger.debug(f"邮件{uid}时间不在范围内，跳过")
                        continue

                    # 解析主题
                    subject = client._get_msg_subject(msg)
                    name, student_id = parse_subject_pattern(subject)
                    
                    # 主题格式校验
                    if not student_id or not name:
                        rejected.append({
                            "学号": "未知",
                            "姓名": "未知",
                            "原主题": subject[:100],  # 截断过长主题
                            "原因": "主题格式错误（示例：薛孜324011234书法班报名申请）"
                        })
                        continue

                    # 黑名单过滤
                    if student_id in blacklist:
                        rejected.append({"学号": student_id, "姓名": name, "原因": "黑名单用户"})
                        continue

                    # 新鸿基直接录取
                    if student_id in new_hongji:
                        admitted.append({"学号": student_id, "姓名": name, "备注": "新鸿基"})
                        continue

                    # 去年已录取
                    if student_id in last_year:
                        rejected.append({"学号": student_id, "姓名": name, "原因": "去年已录取"})
                        continue

                    # 处理附件（捕获附件处理异常）
                    try:
                        attachments = email_processor.save_attachments(msg, student_id, name)
                        docx_files = [a for a in attachments if a.suffix.lower() == ".docx"]
                    except Exception as e:
                        logger.error(f"邮件{uid}附件处理失败: {e}", exc_info=True)
                        rejected.append({"学号": student_id, "姓名": name, "原因": f"附件处理失败: {str(e)[:50]}"})
                        continue

                    # 附件校验
                    if not docx_files:
                        rejected.append({"学号": student_id, "姓名": name, "原因": "缺少DOCX格式申请附件"})
                        continue

                    # 解析DOCX（捕获DOCX解析异常）
                    try:
                        docx_info = parse_docx(str(docx_files[0]))
                    except Exception as e:
                        logger.error(f"邮件{uid}DOCX解析失败: {e}", exc_info=True)
                        rejected.append({"学号": student_id, "姓名": name, "原因": f"申请材料解析失败: {str(e)[:50]}"})
                        continue

                    # DOCX内容校验
                    if not docx_info["is_supported"]:
                        rejected.append({"学号": student_id, "姓名": name, "原因": "非学生资助对象，不符合申请条件"})
                    elif docx_info["reason_length"] < 95:
                        rejected.append({"学号": student_id, "姓名": name, "原因": f"申请理由字数不足（{docx_info['reason_length']}字，需≥95字）"})
                    else:
                        candidates.append((student_id, name, recv_date))

                except Exception as e:
                    # 单封邮件处理失败，计数+1，继续下一封
                    error_count += 1
                    logger.error(f"处理邮件{uid}失败（已跳过）: {str(e)}", exc_info=True)
                    rejected.append({"学号": "未知", "姓名": "未知", "原因": f"邮件处理异常: {str(e)[:50]}"})
                    continue

        # 5. 处理候补名单（即使候选人为空也不报错）
        logger.info(f"邮件处理完成：总计{email_count}封，错误{error_count}封，有效候选{len(candidates)}人")
        remaining_quota = ADMISSION_QUOTA - len(admitted)
        logger.info(f"新鸿基录取{len(admitted)}人，剩余名额{remaining_quota}")
        
        if remaining_quota > 0 and candidates:
            candidates.sort(key=lambda x: x[2])
            admit_candidates = candidates[:remaining_quota]
            reject_candidates = candidates[remaining_quota:]
            
            for sid, name, _ in admit_candidates:
                admitted.append({"学号": sid, "姓名": name, "备注": "非新鸿基（候补）"})
            for sid, name, _ in reject_candidates:
                rejected.append({"学号": sid, "姓名": name, "原因": "符合条件但名额已满"})
        elif candidates:
            for sid, name, _ in candidates:
                rejected.append({"学号": sid, "姓名": name, "原因": "符合条件但名额已满"})

        # 6. 保存结果（最后一步，即使前面有部分错误也保存已有结果）
        try:
            save_results(admitted, rejected)
            logger.info(f"最终结果：录取{len(admitted)}人，拒绝{len(rejected)}人，结果已保存")
        except Exception as e:
            logger.error(f"保存结果失败: {e}", exc_info=True)
            raise RuntimeError(f"筛选完成但结果保存失败: {str(e)}") from e

        logger.info("="*50 + " 书法班筛选流程执行完成 " + "="*50)

    # 捕获主流程致命异常（无法继续执行的错误）
    except RuntimeError as e:
        logger.critical(f"主流程致命错误: {e}", exc_info=True)
        # 兜底：尝试保存已有结果
        if admitted or rejected:
            try:
                save_results(admitted, rejected)
                logger.warning(f"已保存部分结果（录取{len(admitted)}人，拒绝{len(rejected)}人）")
            except:
                pass
        raise  # 重新抛出，让前端捕获返回码
    except Exception as e:
        logger.critical(f"未预期的全局异常: {e}", exc_info=True)
        # 兜底保存
        if admitted or rejected:
            try:
                save_results(admitted, rejected)
            except:
                pass
        raise RuntimeError(f"程序执行异常: {str(e)}") from e

if __name__ == "__main__":
    main()
