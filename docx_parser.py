import re
import logging
from docx import Document
from typing import Dict, Any

logger = logging.getLogger(__name__)

class DocxParser:
    """DOCX申请表解析类"""
    
    @staticmethod
    def parse(filepath: str) -> Dict[str, Any]:
        """解析DOCX中的资助对象和申请理由"""
        result = {
            "is_supported": False,
            "reason_length": 0
        }
        try:
            doc = Document(filepath)
            table = doc.tables[0] if doc.tables else None
            if not table:
                return result
            
            for row in table.rows:
                cells = row.cells
                for cell_idx, cell in enumerate(cells):
                    cell_text = cell.text.strip()
                    # 检测资助对象
                    if "是否为学生资助对象" in cell_text:
                        if cell_idx + 1 < len(cells):
                            next_cell = cells[cell_idx+1].text.strip()
                            result["is_supported"] = (
                                ("是" in next_cell or "为" in next_cell) 
                                and ("不是" not in next_cell)
                            )
                    # 检测申请理由
                    if "申请理由" in cell_text:
                        reason_text = "\n".join([
                            p.text.strip() for p in cell.paragraphs if p.text.strip()
                        ])
                        reason_text = re.sub(r"\s+", "", reason_text)
                        result["reason_length"] = len(reason_text)
            return result
        except Exception as e:
            logger.error(f"解析DOCX失败: {str(e)}")
            return result
